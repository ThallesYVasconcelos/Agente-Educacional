"""
Exportação de conteúdo Markdown para DOCX (Word).

Converte corretamente:
- **negrito** e *itálico* → estilos Word reais (sem asteriscos visíveis)
- # ## ### → estilos Heading 1/2/3
- - bullet / 1. numerado → listas Word nativas
- > blockquote → parágrafo recuado com borda lateral
- --- → separador horizontal
- UTF-8 completo (acentos, ç, ã, etc.)
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


# ---------------------------------------------------------------------------
# Cores da identidade visual
# ---------------------------------------------------------------------------

PRIMARY = RGBColor(46, 91, 255)
DARK = RGBColor(26, 31, 54)
GRAY = RGBColor(90, 99, 120)


# ---------------------------------------------------------------------------
# Parser de spans inline (negrito / itálico)
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"\*\*\*(?P<bi>.+?)\*\*\*"
    r"|\*\*(?P<b>.+?)\*\*"
    r"|__(?P<b2>.+?)__"
    r"|\*(?P<i>.+?)\*"
    r"|_(?P<i2>.+?)_",
    re.DOTALL,
)


def _parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    """Retorna lista de (texto, bold, italic)."""
    spans: list[tuple[str, bool, bool]] = []
    cursor = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > cursor:
            spans.append((text[cursor:m.start()], False, False))
        if m.group("bi"):
            spans.append((m.group("bi"), True, True))
        elif m.group("b") or m.group("b2"):
            spans.append((m.group("b") or m.group("b2"), True, False))
        elif m.group("i") or m.group("i2"):
            spans.append((m.group("i") or m.group("i2"), False, True))
        cursor = m.end()
    if cursor < len(text):
        spans.append((text[cursor:], False, False))
    return spans or [(text, False, False)]


def _strip_md(text: str) -> str:
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,2}(.+?)_{1,2}", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"(?<!\w)[*_]+(?!\w)", "", text)
    return text.strip()


def _clean_llm_output(text: str) -> str:
    text = re.sub(r"```[a-zA-Z]*\n?", "", text)
    text = re.sub(r"```", "", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Helpers DOCX
# ---------------------------------------------------------------------------

def _add_border_left(paragraph, color: str = "2E5BFF", width: int = 12):
    """Adiciona borda esquerda a um parágrafo (estilo blockquote)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_inline_spans(paragraph, spans: list[tuple[str, bool, bool]],
                      base_size: int = 11, color: RGBColor | None = None):
    """Adiciona runs com bold/italic ao parágrafo."""
    for text, bold, italic in spans:
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(base_size)
        if color:
            run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int):
    clean = _strip_md(text)
    p = doc.add_heading(clean, level=level)
    # Aplica cor primária ao título
    for run in p.runs:
        run.font.color.rgb = PRIMARY
    return p


def _bullet_para(doc: Document, spans: list[tuple[str, bool, bool]], level: int = 0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    _add_inline_spans(p, spans)
    return p


def _numbered_para(doc: Document, spans: list[tuple[str, bool, bool]]):
    p = doc.add_paragraph(style="List Number")
    _add_inline_spans(p, spans)
    return p


def _blockquote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    _add_border_left(p)
    run = p.add_run(_strip_md(text))
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p


def _horizontal_rule(doc: Document):
    """Adiciona uma linha horizontal via borda inferior de parágrafo vazio."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8D2F0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------------------
# Configuração de estilos do documento
# ---------------------------------------------------------------------------

def _setup_styles(doc: Document):
    """Ajusta estilos globais do documento."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # Espaçamento entre parágrafos
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = PRIMARY
        h_style.font.size = Pt([18, 14, 12][level - 1])
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(4)


def _add_title_block(doc: Document, title: str):
    """Adiciona o cabeçalho visual do documento."""
    # Faixa de título
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(_strip_md(title))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = PRIMARY
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linha decorativa
    _horizontal_rule(doc)

    # Subtítulo/crédito
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Assistente do Professor — EduRAG (BNCC/PCN)")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = GRAY
    sub_run.italic = True

    doc.add_paragraph()  # espaço


# ---------------------------------------------------------------------------
# Parser de linhas Markdown → documento DOCX
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)")
_BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.*)")
_NUMBERED_RE = re.compile(r"^\s*(\d+)[.)]\s+(.*)")
_RULE_RE = re.compile(r"^\s*[-*_]{3,}\s*$")
_BQ_RE = re.compile(r"^>\s?(.*)")


def markdown_to_docx(markdown_text: str, title: str = "Documento") -> bytes:
    """
    Converte texto Markdown para bytes de arquivo DOCX (Word).

    - UTF-8 completo
    - **negrito** e *itálico* aplicados como estilos Word reais
    - Sem asteriscos visíveis
    """
    markdown_text = _clean_llm_output(markdown_text)

    doc = Document()
    _setup_styles(doc)

    # Margens
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    _add_title_block(doc, title)

    lines = markdown_text.splitlines()

    for raw in lines:
        stripped = raw.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        if _RULE_RE.match(stripped):
            _horizontal_rule(doc)
            continue

        m_bq = _BQ_RE.match(stripped)
        if m_bq:
            _blockquote(doc, m_bq.group(1))
            continue

        m_h = _HEADING_RE.match(stripped)
        if m_h:
            _heading(doc, m_h.group(2).strip(), len(m_h.group(1)))
            continue

        m_b = _BULLET_RE.match(raw)
        if m_b:
            level = len(m_b.group(1)) // 2
            _bullet_para(doc, _parse_inline(m_b.group(2).strip()), level)
            continue

        m_n = _NUMBERED_RE.match(stripped)
        if m_n:
            _numbered_para(doc, _parse_inline(m_n.group(2).strip()))
            continue

        # Parágrafo normal
        p = doc.add_paragraph()
        _add_inline_spans(p, _parse_inline(stripped))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
